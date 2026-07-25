import os
import re
import subprocess
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def create_research_paper_docx():
    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    md_path = os.path.join(base_dir, "Predictive_RL_DVFS_Research_Paper.md")
    docx_path = os.path.join(base_dir, "Predictive_RL_DVFS_Research_Paper.docx")
    pdf_path = os.path.join(base_dir, "Predictive_RL_DVFS_Research_Paper.pdf")
    html_path = os.path.join(base_dir, "IEEE_Predictive_RL_DVFS_Research_Paper.html")

    with open(md_path, "r", encoding="utf-8") as f:
        md_lines = f.readlines()

    doc = docx.Document()

    # Set Margins (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Base Styles
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
    normal_style.paragraph_format.line_spacing = 1.15
    normal_style.paragraph_format.space_after = Pt(6)

    in_table = False
    table_lines = []
    in_code = False
    code_lines = []

    def format_inline_formatting(p, text):
        # Parses markdown bold **text**, italic *text*, code `text`
        parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*'):
                run = p.add_run(part[1:-1])
                run.italic = True
            elif part.startswith('`') and part.endswith('`'):
                run = p.add_run(part[1:-1])
                run.font.name = 'Consolas'
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
            else:
                p.add_run(part)

    def render_table(lines):
        if len(lines) < 2:
            return
        headers = [c.strip() for c in lines[0].strip('|').split('|')]
        rows = []
        for line in lines[2:]:
            if '|' in line:
                rows.append([c.strip() for c in line.strip('|').split('|')])
        
        table = doc.add_table(rows=len(rows)+1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        # Header Row
        hdr_cells = table.rows[0].cells
        for i, header_text in enumerate(headers):
            hdr_cells[i].text = header_text
            set_cell_background(hdr_cells[i], "0B2545")
            set_cell_margins(hdr_cells[i], top=120, bottom=120, left=150, right=150)
            p = hdr_cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.size = Pt(9.5)

        # Data Rows
        for r_idx, row_data in enumerate(rows):
            row_cells = table.rows[r_idx+1].cells
            bg_color = "F8FAFC" if r_idx % 2 == 1 else "FFFFFF"
            for c_idx, cell_value in enumerate(row_data):
                if c_idx < len(row_cells):
                    set_cell_background(row_cells[c_idx], bg_color)
                    set_cell_margins(row_cells[c_idx], top=100, bottom=100, left=150, right=150)
                    p = row_cells[c_idx].paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p.paragraph_format.space_after = Pt(2)
                    format_inline_formatting(p, cell_value)

        doc.add_paragraph() # Spacer

    for line in md_lines:
        raw_line = line.strip()

        # Handle Code Block
        if raw_line.startswith("```"):
            if in_code:
                # Add code box
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.right_indent = Inches(0.25)
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(8)
                run = p.add_run("\n".join(code_lines))
                run.font.name = 'Consolas'
                run.font.size = Pt(9.0)
                run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
                code_lines = []
                in_code = False
            else:
                in_code = True
                code_lines = []
            continue

        if in_code:
            code_lines.append(line.rstrip())
            continue

        # Handle Table
        if raw_line.startswith("|") and raw_line.endswith("|"):
            if not in_table:
                in_table = True
                table_lines = [raw_line]
            else:
                table_lines.append(raw_line)
            continue
        else:
            if in_table:
                render_table(table_lines)
                table_lines = []
                in_table = False

        if not raw_line:
            continue

        # Title (# )
        if raw_line.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            run = p.add_run(raw_line[2:].strip())
            run.font.size = Pt(20)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x0B, 0x25, 0x45)

        # Section Heading (## )
        elif raw_line.startswith("## "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(16)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(raw_line[3:].strip())
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x0B, 0x25, 0x45)

        # Sub-heading (### )
        elif raw_line.startswith("### "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(raw_line[4:].strip())
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x13, 0x40, 0x74)

        # Divider (---)
        elif raw_line == "---":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(10)
            run = p.add_run("_________________________________________________________________________________")
            run.font.color.rgb = RGBColor(0xCB, 0xD5, 0xE1)
            run.font.size = Pt(8)

        # Author / Metadata block
        elif raw_line.startswith("**Author:**") or raw_line.startswith("**Affiliation:**"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(2)
            format_inline_formatting(p, raw_line)

        # Standard Paragraph
        else:
            p = doc.add_paragraph()
            format_inline_formatting(p, raw_line)

    if in_table:
        render_table(table_lines)

    try:
        doc.save(docx_path)
        print(f"Successfully generated DOCX document at: {docx_path}")
    except PermissionError:
        alt_docx_path = os.path.join(base_dir, "Predictive_RL_DVFS_Research_Paper_Updated.docx")
        doc.save(alt_docx_path)
        print(f"Original DOCX file locked by another application. Saved to: {alt_docx_path}")

    # Try exporting PDF using Headless Edge / Chrome
    edge_paths = [
        r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe",
        r"C:\Program Files\Microsoft\Edge\Application\msedge.exe",
        r"C:\Program Files\Google\Chrome\Application\chrome.exe"
    ]

    edge_bin = None
    for path in edge_paths:
        if os.path.exists(path):
            edge_bin = path
            break

    if edge_bin:
        cmd = [
            edge_bin,
            "--headless",
            "--disable-gpu",
            f"--print-to-pdf={pdf_path}",
            html_path
        ]
        try:
            subprocess.run(cmd, check=True)
            print(f"Successfully generated PDF document at: {pdf_path}")
        except Exception as e:
            print(f"PDF generation via browser failed: {e}")

if __name__ == "__main__":
    create_research_paper_docx()
